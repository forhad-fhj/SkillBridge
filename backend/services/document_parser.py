import PyPDF2
import pdfplumber
from docx import Document
from typing import Optional
import io

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from PDF file using PyPDF2 and pdfplumber as fallback.
    
    Args:
        file_bytes: PDF file content as bytes
        
    Returns:
        Extracted text content
    """
    text = ""
    
    try:
        # Try PyPDF2 first (faster)
        pdf_file = io.BytesIO(file_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        
        # If PyPDF2 didn't extract much text, try pdfplumber
        if len(text.strip()) < 100:
            pdf_file.seek(0)
            with pdfplumber.open(pdf_file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
    
    except Exception as e:
        raise ValueError(f"Error extracting text from PDF: {str(e)}")
    
    return text.strip()

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from DOCX file.
    
    Args:
        file_bytes: DOCX file content as bytes
        
    Returns:
        Extracted text content
    """
    try:
        docx_file = io.BytesIO(file_bytes)
        doc = Document(docx_file)
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip()
    
    except Exception as e:
        raise ValueError(f"Error extracting text from DOCX: {str(e)}")

def parse_document(file_bytes: bytes, filename: str) -> str:
    """
    Parse document and extract text based on file type.
    
    Args:
        file_bytes: File content as bytes
        filename: Original filename with extension
        
    Returns:
        Extracted text content
    """
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.pdf'):
        return extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith('.docx'):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {filename}. Only PDF and DOCX are supported.")

def clean_text(text: str) -> str:
    """
    Clean and normalize extracted text.
    
    Args:
        text: Raw extracted text
        
    Returns:
        Cleaned text
    """
    # Remove excessive whitespace
    text = " ".join(text.split())
    
    # Remove special characters but keep punctuation
    # text = re.sub(r'[^\w\s\.\,\-\(\)]', '', text)
    
    return text.strip()
